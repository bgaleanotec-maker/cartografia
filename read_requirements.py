from docx import Document
import os

def extract_to_file():
    try:
        file_path = "req.docx"
        doc = Document(file_path)
        
        with open("requirements.txt", "w", encoding="utf-8") as f:
            f.write("--- PARAGRAPHS ---\n")
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    f.write(text + "\n")
                    
            f.write("\n--- TABLES ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_parts = []
                    for cell in row.cells:
                        txt = cell.text.strip()
                        # simple clean
                        txt = " ".join(txt.split())
                        row_parts.append(txt)
                    f.write(" | ".join(row_parts) + "\n")
                f.write("\n")
                
        print("Done.")
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    extract_to_file()
